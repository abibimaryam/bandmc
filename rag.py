# import os
# import re
# import pandas as pd
# from langchain_chroma import Chroma
# from langchain.prompts import ChatPromptTemplate
# from langchain_community.llms.ollama import Ollama
# from embedding_functiom import embedding_function
# from config import CHROMA_PATH, model
# from ddgs import DDGS
# from docx import Document
# import pdfplumber
# import time


# SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".docx", ".csv", ".xlsx"]

# PROMPT_TEMPLATE = """
# Ниже приведён контекст из медицинской базы данных и интернета:

# {context}

# Пользователь сделал следующее утверждение:

# {question}

# В утверждении указаны следующие ключевые элементы: {key_entities}

# Проанализируй утверждение в контексте и найди ВСЕ несоответствия или ошибки, особенно проверь, совпадают ли производитель и страна-производитель с базой данных и интернет-источниками.

# Выведи список только тех пунктов, где утверждение НЕ СООТВЕТСТВУЕТ данным из контекста, с кратким объяснением.

# Если несоответствий нет, ответь "Несоответствий не обнаружено."
# """

# # --- Веб-поиск
# def search_tool(query: str) -> str:
#     with DDGS() as ddgs:
#         results = [r["body"] for r in ddgs.text(query, max_results=3)]
#     return "\n\n".join(results)


# def extract_key_entities(user_input: str) -> str:
#     prompt = f"""
# Извлеки ключевые элементы из следующего запроса, такие как:

# - торговое название лекарства,
# - международное непатентованное название (если есть),

# Верни их через запятую — одной строкой. Если ничего не найдено, верни "None".

# Запрос: "{user_input}"
# """
#     response = model.invoke(prompt).strip()
#     print(f"🔑 Извлечённые ключевые элементы: {response}")
#     return response if response.lower() != "none" else None


# def query_rag_with_web_search(user_input: str):
#     key_phrases = extract_key_entities(user_input)
#     query_text = key_phrases or user_input
#     filter_name = key_phrases.split(",")[0].strip().lower() if key_phrases else None

#     db = Chroma(
#         collection_name="meds",
#         persist_directory=CHROMA_PATH,
#         embedding_function=embedding_function()
#     )
#     results = db.similarity_search_with_score(query_text, k=20)

#     if filter_name:
#         results = [(doc, score) for doc, score in results
#                    if filter_name in doc.metadata.get("торговое_название", "").lower()]

#     local_context = "\n\n---\n\n".join([doc.page_content for doc, _ in results]) if results else ""
#     web_results = search_tool(user_input)

#     combined_context = f"Локальная база данных:\n{local_context}\n\nИнтернет данные:\n{web_results}"

#     prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
#     prompt = prompt_template.format(
#         context=combined_context,
#         question=user_input,
#         key_entities=key_phrases or "None"
#     )

#     response_text = model.invoke(prompt).strip()
#     return response_text


# # --- Функции чтения файлов
# def read_txt(filepath: str) -> str:
#     with open(filepath, "r", encoding="utf-8") as f:
#         return f.read()


# def read_docx(filepath: str) -> str:
#     doc = Document(filepath)
#     return "\n".join([p.text for p in doc.paragraphs])


# def read_pdf(filepath: str) -> str:
#     text = ""
#     with pdfplumber.open(filepath) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() + "\n"
#     return text


# def read_excel_csv(filepath: str) -> str:
#     ext = os.path.splitext(filepath)[-1].lower()
#     df = pd.read_excel(filepath, header=0) if ext == ".xlsx" else pd.read_csv(filepath, header=0)

#     print("📋 Заголовки таблицы:", df.columns.tolist())

#     # Добавим заголовки к каждой строке
#     rows = []
#     for _, row in df.iterrows():
#         row_str = ", ".join([f"{col.strip()}: {str(val).strip()}" for col, val in row.items()])
#         rows.append(row_str)

#     return "\n".join(rows)



# # --- Универсальный загрузчик текста
# def extract_text_from_file(filepath: str) -> str:
#     ext = os.path.splitext(filepath)[-1].lower()
#     if ext == ".txt":
#         return read_txt(filepath)
#     elif ext == ".docx":
#         return read_docx(filepath)
#     elif ext == ".pdf":
#         return read_pdf(filepath)
#     elif ext in [".csv", ".xlsx"]:
#         return read_excel_csv(filepath)
#     else:
#         raise ValueError(f"❌ Формат {ext} не поддерживается. Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS)}")


# def split_document_into_statements(text: str) -> list:
#     # Разделение по строкам или по предложениям
#     return [s.strip() for s in re.split(r'\n+|(?<=[.!?])\s+', text) if s.strip()]




# def query_with_retries(statement: str, retries: int = 2, delay: float = 2.0) -> str:
#     for attempt in range(1, retries + 1):
#         try:
#             return query_rag_with_web_search(statement)
#         except Exception as e:
#             print(f"⚠️ Ошибка при проверке утверждения (попытка {attempt}/{retries}): {e}")
#             if attempt < retries:
#                 print(f"⏳ Ждём {delay} сек и повторяем...")
#                 time.sleep(delay)
#             else:
#                 return f"⚠️ Ошибка после {retries} попыток: {e}"

# # --- Главная функция обработки документа
# def process_document(filepath: str):
#     if not os.path.exists(filepath):
#         print(f"❌ Файл {filepath} не найден.")
#         return

#     print(f"📂 Загрузка файла: {filepath}")
#     try:
#         text = extract_text_from_file(filepath)
#     except Exception as e:
#         print(f"Ошибка при чтении файла: {e}")
#         return

#     statements = split_document_into_statements(text)
#     print(f"📄 Найдено {len(statements)} утверждений.\n")

#     all_results = []
#     for i, statement in enumerate(statements, start=1):
#         print(f"\n🔍 Утверждение {i}:\n{statement}")
#         result = query_with_retries(statement, retries=2, delay=2)
#         all_results.append((statement, result))

#     print("\n🧾 Итоговый отчёт по несоответствиям:\n")
#     for i, (stmt, mismatch) in enumerate(all_results, start=1):
#         print(f"{i}. 📌 {stmt}")
#         print(f"   ❗ Проверка: {mismatch}")
#         print("------")

# # --- Точка входа
# if __name__ == "__main__":
#     FILEPATH = "/home/maryam/Документы/projects/bandmc/test_data/test.xlsx"  # путь к файлу
#     process_document(FILEPATH)


# import os
# import re
# import time
# import pandas as pd
# from docx import Document
# import pdfplumber
# from ddgs import DDGS
# from langchain_chroma import Chroma
# from langchain.prompts import ChatPromptTemplate
# from langchain_community.llms.ollama import Ollama
# from embedding_functiom import embedding_function
# from config import CHROMA_PATH, model

# # SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".docx", ".csv", ".xlsx"]

# # PROMPT_TEMPLATE = """
# # Ниже приведён контекст из медицинской базы данных и интернета:

# # {context}

# # Пользователь сделал следующее утверждение:

# # {question}

# # В утверждении указаны следующие ключевые элементы: {key_entities}

# # Проанализируй утверждение в контексте и найди ВСЕ несоответствия или ошибки.

# # Особенно внимательно проверь:
# # - Совпадает ли торговое наименование;
# # - Совпадает ли международное непатентованное название;
# # - Совпадает ли производитель;
# # - Совпадает ли страна-производитель.

# # Если какое-либо из этих значений отличается — обязательно зафиксируй это.

# # Выведи список только тех пунктов, где утверждение НЕ СООТВЕТСТВУЕТ данным из контекста, с кратким объяснением.

# # Если несоответствий нет, ответь "Несоответствий не обнаружено."
# # """

# # # --- Веб-поиск
# # def search_tool(query: str, retries=3, delay=2) -> str:
# #     for attempt in range(retries):
# #         try:
# #             with DDGS() as ddgs:
# #                 results = [r["body"] for r in ddgs.text(query, max_results=3)]
# #             return "\n\n".join(results)
# #         except Exception as e:
# #             print(f"⚠️ Ошибка веб-поиска: {e}")
# #             if attempt < retries - 1:
# #                 print(f"⏳ Повтор через {delay} сек...")
# #                 time.sleep(delay)
# #     return "❌ Не удалось получить данные из интернета."

# # # --- Извлечение ключевых элементов
# # def extract_key_entities(user_input: str) -> str:
# #     prompt = f"""
# # Извлеки ключевые элементы из следующего запроса, такие как:

# # - торговое название лекарства,
# # - международное непатентованное название (если есть),

# # Верни их через запятую — одной строкой. Если ничего не найдено, верни "None".

# # Запрос: "{user_input}"
# # """
# #     response = model.invoke(prompt).strip()
# #     print(f"🔑 Извлечённые ключевые элементы: {response}")
# #     return response if response.lower() != "none" else None

# # # --- Поиск несоответствий
# # def query_rag_with_web_search(user_input: str):
# #     key_phrases = extract_key_entities(user_input)
# #     query_text = key_phrases or user_input
# #     filter_name = key_phrases.split(",")[0].strip().lower() if key_phrases else None

# #     db = Chroma(
# #         collection_name="meds",
# #         persist_directory=CHROMA_PATH,
# #         embedding_function=embedding_function()
# #     )
# #     results = db.similarity_search_with_score(query_text, k=20)

# #     if filter_name:
# #         results = [(doc, score) for doc, score in results
# #                    if filter_name in doc.metadata.get("торговое_название", "").lower()]

# #     if not results:
# #         print(f"⚠️ Не найдено документов по фильтру '{filter_name}'")
# #         local_context = ""
# #     else:
# #         print(f"✅ Найдено {len(results)} документов по фильтру '{filter_name}'")
# #         print("🔍 Пример документа:")
# #         doc = results[0][0]
# #         print(doc.page_content[:500])
# #         print("📌 Метаданные:", doc.metadata)
# #         local_context = "\n\n---\n\n".join([doc.page_content for doc, _ in results])

# #     print("🌐 Выполняется веб-поиск...")
# #     web_results = search_tool(user_input)
# #     print("🌐 Web-контекст:\n", web_results[:300])

# #     combined_context = f"Локальная база данных:\n{local_context}\n\nИнтернет данные:\n{web_results}"

# #     prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
# #     prompt = prompt_template.format(
# #         context=combined_context,
# #         question=user_input,
# #         key_entities=key_phrases or "None"
# #     )

# #     response_text = model.invoke(prompt).strip()
# #     return response_text

# # # --- Чтение файлов
# # def read_txt(filepath: str) -> str:
# #     with open(filepath, "r", encoding="utf-8") as f:
# #         return f.read()

# # def read_docx(filepath: str) -> str:
# #     doc = Document(filepath)
# #     return "\n".join([p.text for p in doc.paragraphs])

# # def read_pdf(filepath: str) -> str:
# #     text = ""
# #     with pdfplumber.open(filepath) as pdf:
# #         for page in pdf.pages:
# #             page_text = page.extract_text()
# #             if page_text:
# #                 text += page_text + "\n"
# #     return text

# # def read_excel_csv(filepath: str) -> str:
# #     ext = os.path.splitext(filepath)[-1].lower()
# #     df = pd.read_excel(filepath, header=0) if ext == ".xlsx" else pd.read_csv(filepath, header=0)
# #     rows = []
# #     for _, row in df.iterrows():
# #         row_str = ", ".join([f"{col.strip()}: {str(val).strip()}" for col, val in row.items()])
# #         rows.append(row_str)
# #     return "\n".join(rows)

# # def extract_text_from_file(filepath: str) -> str:
# #     ext = os.path.splitext(filepath)[-1].lower()
# #     if ext == ".txt":
# #         return read_txt(filepath)
# #     elif ext == ".docx":
# #         return read_docx(filepath)
# #     elif ext == ".pdf":
# #         return read_pdf(filepath)
# #     elif ext in [".csv", ".xlsx"]:
# #         return read_excel_csv(filepath)
# #     else:
# #         raise ValueError(f"❌ Формат {ext} не поддерживается. Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS)}")

# # def split_document_into_statements(text: str) -> list:
# #     return [s.strip() for s in re.split(r'\n+|(?<=[.!?])\s+', text) if s.strip()]

# # # --- Попытки с повторами
# # def query_with_retries(statement: str, retries: int = 2, delay: float = 2.0) -> str:
# #     for attempt in range(1, retries + 1):
# #         try:
# #             return query_rag_with_web_search(statement)
# #         except Exception as e:
# #             print(f"⚠️ Ошибка при проверке (попытка {attempt}/{retries}): {e}")
# #             if attempt < retries:
# #                 print(f"⏳ Повтор через {delay} сек...")
# #                 time.sleep(delay)
# #     return f"❌ Не удалось проверить утверждение после {retries} попыток."

# # # --- Обработка документа
# # def process_document(filepath: str):
# #     if not os.path.exists(filepath):
# #         print(f"❌ Файл {filepath} не найден.")
# #         return

# #     print(f"📂 Загрузка файла: {filepath}")
# #     try:
# #         text = extract_text_from_file(filepath)
# #     except Exception as e:
# #         print(f"❌ Ошибка при чтении файла: {e}")
# #         return

# #     statements = split_document_into_statements(text)
# #     print(f"📄 Найдено {len(statements)} утверждений.\n")

# #     all_results = []
# #     for i, statement in enumerate(statements, start=1):
# #         print(f"\n🔍 Утверждение {i}:\n{statement}")
# #         result = query_with_retries(statement)
# #         all_results.append((statement, result))

# #     print("\n🧾 Итоговый отчёт по несоответствиям:\n")
# #     for i, (stmt, mismatch) in enumerate(all_results, start=1):
# #         print(f"{i}. 📌 {stmt}")
# #         print(f"   ❗ Проверка: {mismatch}")
# #         print("------")

# # # --- Точка входа
# # if __name__ == "__main__":
# #     FILEPATH = "/home/maryam/Документы/projects/bandmc/test_data/test.xlsx"  # 🔁 Подставь свой путь
# #     process_document(FILEPATH)

import os
import re
import time
import pandas as pd
from docx import Document
import pdfplumber
from ddgs import DDGS
from langchain_chroma import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_community.llms.ollama import Ollama
from embedding_functiom import embedding_function
from config import CHROMA_PATH, model

SUPPORTED_EXTENSIONS = [".txt", ".pdf", ".docx", ".csv", ".xlsx"]

PROMPT_TEMPLATE = """
Ниже приведён контекст из медицинской базы данных и интернета:

{context}

Пользователь сделал следующее утверждение:

{question}

В утверждении указаны следующие ключевые элементы: {key_entities}

Проанализируй утверждение в контексте и найди ВСЕ несоответствия или ошибки.

Особенно внимательно проверь:
- Совпадает ли торговое наименование;
- Совпадает ли международное непатентованное название;
- Совпадает ли производитель;
- Совпадает ли страна-производитель.

Если какое-либо из этих значений отличается — обязательно зафиксируй это.

Выведи список только тех пунктов, где утверждение НЕ СООТВЕТСТВУЕТ данным из контекста, с кратким объяснением.

Если несоответствий нет, ответь "Несоответствий не обнаружено."
"""

# --- Веб-поиск
def search_tool(query: str, retries=3, delay=2) -> str:
    for attempt in range(retries):
        try:
            with DDGS() as ddgs:
                results = [r["body"] for r in ddgs.text(query, max_results=3)]
            return "\n\n".join(results)
        except Exception as e:
            print(f"⚠️ Ошибка веб-поиска: {e}")
            if attempt < retries - 1:
                print(f"⏳ Повтор через {delay} сек...")
                time.sleep(delay)
    return "❌ Не удалось получить данные из интернета."

# --- Извлечение ключевых элементов
def extract_key_entities(user_input: str) -> str:
    prompt = f"""
Извлеки ключевые элементы из следующего запроса, такие как:

- торговое название лекарства,
- международное непатентованное название (если есть),

Верни их через запятую — одной строкой. Если ничего не найдено, верни "None".

Запрос: "{user_input}"
"""
    response = model.invoke(prompt)
    response_text = response.content.strip()
    print(f"🔑 Извлечённые ключевые элементы: {response_text}")
    return response_text if response_text.lower() != "none" else None

# --- Поиск несоответствий
def query_rag_with_web_search(user_input: str):
    key_phrases = extract_key_entities(user_input)
    query_text = key_phrases or user_input
    filter_name = key_phrases.split(",")[0].strip().lower() if key_phrases else None

    db = Chroma(
        collection_name="meds",
        persist_directory=CHROMA_PATH,
        embedding_function=embedding_function()
    )
    results = db.similarity_search_with_score(query_text, k=20)

    if filter_name:
        results = [(doc, score) for doc, score in results
                   if filter_name in doc.metadata.get("торговое_название", "").lower()]

    if not results:
        print(f"⚠️ Не найдено документов по фильтру '{filter_name}'")
        local_context = ""
    else:
        print(f"✅ Найдено {len(results)} документов по фильтру '{filter_name}'")
        print("🔍 Пример документа:")
        doc = results[0][0]
        print(doc.page_content[:500])
        print("📌 Метаданные:", doc.metadata)
        local_context = "\n\n---\n\n".join([doc.page_content for doc, _ in results])

    print("🌐 Выполняется веб-поиск...")
    web_results = search_tool(user_input)
    print("🌐 Web-контекст:\n", web_results[:300])

    combined_context = f"Локальная база данных:\n{local_context}\n\nИнтернет данные:\n{web_results}"

    prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
    prompt = prompt_template.format(
        context=combined_context,
        question=user_input,
        key_entities=key_phrases or "None"
    )

    response = model.invoke(prompt)
    response_text = response.content.strip()
    return response_text

# --- Чтение файлов
def read_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()

def read_docx(filepath: str) -> str:
    doc = Document(filepath)
    return "\n".join([p.text for p in doc.paragraphs])

def read_pdf(filepath: str) -> str:
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def read_excel_csv(filepath: str) -> str:
    ext = os.path.splitext(filepath)[-1].lower()
    df = pd.read_excel(filepath, header=0) if ext == ".xlsx" else pd.read_csv(filepath, header=0)
    rows = []
    for _, row in df.iterrows():
        row_str = ", ".join([f"{col.strip()}: {str(val).strip()}" for col, val in row.items()])
        rows.append(row_str)
    return "\n".join(rows)

def extract_text_from_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[-1].lower()
    if ext == ".txt":
        return read_txt(filepath)
    elif ext == ".docx":
        return read_docx(filepath)
    elif ext == ".pdf":
        return read_pdf(filepath)
    elif ext in [".csv", ".xlsx"]:
        return read_excel_csv(filepath)
    else:
        raise ValueError(f"❌ Формат {ext} не поддерживается. Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS)}")

def split_document_into_statements(text: str) -> list:
    return [s.strip() for s in re.split(r'\n+|(?<=[.!?])\s+', text) if s.strip()]

# --- Попытки с повторами
def query_with_retries(statement: str, retries: int = 2, delay: float = 2.0) -> str:
    for attempt in range(1, retries + 1):
        try:
            return query_rag_with_web_search(statement)
        except Exception as e:
            print(f"⚠️ Ошибка при проверке (попытка {attempt}/{retries}): {e}")
            if attempt < retries:
                print(f"⏳ Повтор через {delay} сек...")
                time.sleep(delay)
    return f"❌ Не удалось проверить утверждение после {retries} попыток."

# --- Обработка документа
def process_document(filepath: str):
    if not os.path.exists(filepath):
        print(f"❌ Файл {filepath} не найден.")
        return

    print(f"📂 Загрузка файла: {filepath}")
    try:
        text = extract_text_from_file(filepath)
    except Exception as e:
        print(f"❌ Ошибка при чтении файла: {e}")
        return

    statements = split_document_into_statements(text)
    print(f"📄 Найдено {len(statements)} утверждений.\n")

    all_results = []
    for i, statement in enumerate(statements, start=1):
        print(f"\n🔍 Утверждение {i}:\n{statement}")
        result = query_with_retries(statement)
        all_results.append((statement, result))

    print("\n🧾 Итоговый отчёт по несоответствиям:\n")
    for i, (stmt, mismatch) in enumerate(all_results, start=1):
        print(f"{i}. 📌 {stmt}")
        print(f"   ❗ Проверка: {mismatch}")
        print("------")

# --- Точка входа
if __name__ == "__main__":
    FILEPATH = "/home/maryam/Документы/projects/bandmc/test_data/test.xlsx"  # 🔁 Подставь свой путь
    process_document(FILEPATH)
